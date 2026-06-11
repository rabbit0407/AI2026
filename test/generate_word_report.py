import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# Paths
project_dir = r"c:\Users\user\Desktop\test"
docx_path = os.path.join(project_dir, "삶의_만족도와_자살률_상관관계_분석_보고서.docx")

doc = Document()

# Page setup - Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Style setup
styles = doc.styles

# Helper function to style text run
def format_run(run, font_name="Malgun Gothic", size_pt=10, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    # Apply font name to East Asia text as well
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# Helper function to set cell background color
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# Helper function to center text in cell vertically and align horizontally
def align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER):
    # Vertical
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), vertical)
    tcPr.append(vAlign)
    # Horizontal
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal

# Color Palette (Dark Blue Theme)
COLOR_PRIMARY = RGBColor(31, 78, 120)     # Deep Blue
COLOR_SECONDARY = RGBColor(89, 89, 89)    # Gray
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_LIGHT_GRAY = "F2F2F2"               # Hex for cell shading
COLOR_PRIMARY_HEX = "1F4E78"              # Hex for header cells

# --- TITLE ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(20)
p_title.paragraph_format.space_after = Pt(6)
run_title = p_title.add_run("삶의 만족도와 자살률의 상관관계 분석 보고서\n(2020 ~ 2024)")
format_run(run_title, size_pt=20, bold=True, color_rgb=COLOR_PRIMARY)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.space_after = Pt(24)
run_subtitle = p_subtitle.add_run("KOSIS 사회조사 및 사망원인통계 데이터 기반 분석")
format_run(run_subtitle, size_pt=12, italic=True, color_rgb=COLOR_SECONDARY)

# --- 1. 요약 및 핵심 통찰 ---
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)
format_run(h1.add_run("1. 요약 및 핵심 통찰 (Executive Summary)"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

bullet_points = [
    ("상관관계의 비유의성 (전체 통합): ", "2020년부터 2024년까지의 5개년 전체 시도별 데이터를 통합하여 분석한 결과, 삶의 만족도 통합 점수와 자살률 간의 단순 선형 상관관계는 통계적으로 유의하지 않은 수준으로 나타났습니다."),
    ("코로나19 시기(2020년)의 음의 상관관계: ", "팬데믹 초기인 2020년에는 삶의 만족도가 높은 지역일수록 자살률이 낮아지는 유의미한 음(-)의 상관관계(Pearson r = -0.50, p < 0.05)가 관찰되었으나, 이후 연도에서는 이러한 관계가 해체되었습니다."),
    ("세부 설문 항목 분석 (보통 응답의 역설): ", "만족 여부를 묻는 세부 비율 분석 결과, 매우 만족하거나 매우 불만족하는 극단적 응답이 높은 지역보다 \"보통(Neutral)\"이라고 답한 비율이 높은 지역일수록 자살률이 증가하는 뚜렷한 양(+)의 상관관계(Pearson r = 0.41)가 관찰되었습니다. 이는 정서적 무관심이나 미온적 상태가 지역 자살 위험과 깊게 연관되어 있을 수 있음을 시사합니다."),
    ("자살률-만족도의 역설 (Daly's Paradox): ", "2024년 시도별 분석 결과, 삶의 만족도가 전국 최상위권인 충청남도와 제주특별자치도에서 오히려 자살률이 각각 2위(34.8명)와 1위(36.3명)를 기록하여 '행복한 지역의 높은 자살률 역설'이 강하게 확인되었습니다."),
    ("성별 불일치 (Gender Disparity): ", "남성과 여성의 삶의 만족도는 거의 동일하나, 자살률은 남성이 여성보다 약 2.5배 높았으며 최근 남성 자살률의 증가세가 극도로 가파릅니다.")
]

for title, body in bullet_points:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run_bold = p.add_run(title)
    format_run(run_bold, size_pt=10, bold=True)
    run_body = p.add_run(body)
    format_run(run_body, size_pt=10)

# --- 2. 데이터 개요 및 지표 산출 ---
h2 = doc.add_paragraph()
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(6)
format_run(h2.add_run("2. 데이터 개요 및 지표 산출"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

p_data = doc.add_paragraph()
p_data.paragraph_format.space_after = Pt(6)
run_data = p_data.add_run(
    "본 분석에서는 통계청 국가통계포털(KOSIS)의 두 가지 원천 통계 데이터를 결합하였습니다.\n"
    "1. 삶의 만족도 데이터: 사회조사 기반의 주관적 정서 지표 (매우 만족, 약간 만족, 보통, 약간 불만족, 매우 불만족 비율)\n"
    "2. 자살률 데이터: 사망원인통계 기반의 인구 10만 명당 자살률\n\n"
    "정량적인 비교 분석을 위해, 사회조사의 5가지 만족도 설문 답변 비율을 반영한 5점 만점의 '삶의 만족도 점수(Satisfaction Score)'를 아래 수식으로 산출하여 활용했습니다."
)
format_run(run_data, size_pt=10)

p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.space_before = Pt(8)
p_formula.paragraph_format.space_after = Pt(12)
run_formula = p_formula.add_run(
    "만족도 점수 = [ (매우 만족 × 5) + (약간 만족 × 4) + (보통 × 3) + (약간 불만족 × 2) + (매우 불만족 × 1) ] / 100"
)
format_run(run_formula, size_pt=11, bold=True, color_rgb=COLOR_PRIMARY)


# --- 3. 전국 평균 추이 및 성별 격차 ---
h3 = doc.add_paragraph()
h3.paragraph_format.space_before = Pt(18)
h3.paragraph_format.space_after = Pt(6)
format_run(h3.add_run("3. 전국 평균 추이 및 성별 격차 (National & Gender Trends)"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

# Table 1: National Trends
national_data = [
    ["성별", "연도", "만족 비율 (%)", "불만족 비율 (%)", "만족도 점수 (5점)", "10만 명당 자살률 (명)"],
    ["전체", "2020", "42.7", "12.5", "3.42", "25.7"],
    ["", "2021", "34.0", "22.9", "3.16", "26.0"],
    ["", "2022", "43.3", "14.1", "3.39", "25.2"],
    ["", "2023", "42.2", "14.5", "3.36", "27.3"],
    ["", "2024", "40.1", "12.7", "3.37", "29.1"],
    ["남성", "2020", "43.5", "11.8", "3.45", "35.5"],
    ["", "2021", "33.8", "23.5", "3.15", "35.9"],
    ["", "2022", "43.4", "14.1", "3.39", "35.3"],
    ["", "2023", "42.5", "14.6", "3.37", "38.3"],
    ["", "2024", "40.5", "12.4", "3.38", "41.8"],
    ["여성", "2020", "41.9", "13.0", "3.40", "15.9"],
    ["", "2021", "34.3", "22.4", "3.17", "16.2"],
    ["", "2022", "43.2", "14.0", "3.39", "15.1"],
    ["", "2023", "42.0", "14.5", "3.36", "16.5"],
    ["", "2024", "39.7", "13.1", "3.35", "16.6"]
]

table1 = doc.add_table(rows=len(national_data), cols=6)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row_data in enumerate(national_data):
    for c_idx, val in enumerate(row_data):
        cell = table1.cell(r_idx, c_idx)
        cell.text = val
        # Format Text
        is_header = (r_idx == 0)
        bold = is_header or (c_idx == 0 and val != "")
        color_rgb = COLOR_BLACK
        run = cell.paragraphs[0].runs[0]
        format_run(run, size_pt=9.5, bold=bold, color_rgb=color_rgb)
        
        # Style Cell
        if is_header:
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            run.font.color.rgb = RGBColor(255, 255, 255) # White text for header
        elif r_idx % 2 == 1:
            set_cell_background(cell, COLOR_LIGHT_GRAY)
            
        align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER)

# Spacing after table
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(12)

# Insert images for section 3
# Image 1: national_trends.png
img_path1 = os.path.join(project_dir, "images", "national_trends.png")
if os.path.exists(img_path1):
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img1.add_run().add_picture(img_path1, width=Inches(5.5))
    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap1.add_run("<그림 1> 연도별 전국 평균 삶의 만족도 및 자살률 추이"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap1.paragraph_format.space_after = Pt(12)

# Image 2: gender_disparity_trends.png
img_path2 = os.path.join(project_dir, "images", "gender_disparity_trends.png")
if os.path.exists(img_path2):
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.add_run().add_picture(img_path2, width=Inches(5.5))
    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap2.add_run("<그림 2> 전국 성별 자살률 및 삶의 만족도 격차 비교"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap2.paragraph_format.space_after = Pt(18)


# --- 4. 상관관계 분석 결과 ---
h4 = doc.add_paragraph()
h4.paragraph_format.space_before = Pt(18)
h4.paragraph_format.space_after = Pt(6)
format_run(h4.add_run("4. 상관관계 분석 결과 (Correlation Analysis)"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

p_corr_desc = doc.add_paragraph()
format_run(p_corr_desc.add_run(
    "전국 통계를 제외하고 17개 시도별로 데이터를 매칭하여 상관관계 통계 분석을 수행했습니다.\n"
    "아래 표 2는 5개년 전체 기간과 전 지역 데이터를 병합(N=85)하여 계산한 전체 통합 상관관계 결과입니다."
), size_pt=10)

# Table 2: Pooled Correlation
pooled_corr_data = [
    ["분석 대상 성별", "독립 변수 (X)", "종속 변수 (Y)", "피어슨 계수 (r)", "피어슨 p-value", "스피어먼 계수 (ρ)", "스피어먼 p-value"],
    ["전체 (Total)", "만족 비율 (%)", "자살률", "-0.1704", "1.19e-01", "-0.1359", "2.15e-01"],
    ["", "불만족 비율 (%)", "자살률", "-0.1767", "1.06e-01", "-0.1685", "1.23e-01"],
    ["", "만족도 점수", "자살률", "-0.0452", "6.81e-01", "-0.0192", "8.61e-01"],
    ["남성 (Male)", "만족 비율 (%)", "자살률", "-0.1208", "2.71e-01", "-0.0158", "8.86e-01"],
    ["", "불만족 비율 (%)", "자살률", "-0.1765", "1.06e-01", "-0.1716", "1.16e-01"],
    ["", "만족도 점수", "자살률", "-0.0304", "7.82e-01", "0.0391", "7.23e-01"],
    ["여성 (Female)", "만족 비율 (%)", "자살률", "-0.0422", "7.02e-01", "-0.0591", "5.91e-01"],
    ["", "불만족 비율 (%)", "자살률", "-0.0674", "5.40e-01", "-0.0450", "6.83e-01"],
    ["", "만족도 점수", "자살률", "0.0136", "9.01e-01", "-0.0180", "8.70e-01"]
]

table2 = doc.add_table(rows=len(pooled_corr_data), cols=7)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row_data in enumerate(pooled_corr_data):
    for c_idx, val in enumerate(row_data):
        cell = table2.cell(r_idx, c_idx)
        cell.text = val
        is_header = (r_idx == 0)
        bold = is_header or (c_idx == 0 and val != "")
        run = cell.paragraphs[0].runs[0]
        format_run(run, size_pt=9, bold=bold)
        if is_header:
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            run.font.color.rgb = RGBColor(255, 255, 255)
        elif r_idx % 2 == 1:
            set_cell_background(cell, COLOR_LIGHT_GRAY)
        align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER)

p_corr_desc2 = doc.add_paragraph()
p_corr_desc2.paragraph_format.space_before = Pt(12)
format_run(p_corr_desc2.add_run(
    "또한, 연도별 시도별로 단년도 단위로 쪼개어 만족도 점수 및 만족 비율과 자살률 간의 상관관계를 분석한 결과는 아래 표 3과 같습니다.\n"
    "2020년에는 통계적으로 유의미한 중간 세기의 음(-)의 상관관계가 관찰되었으나, 그 외 연도에는 상관관계가 매우 낮거나 양의 방향으로 희석되는 복합적인 양상을 보였습니다."
), size_pt=10)

# Table 3: Year-by-Year Correlation
yby_corr_data = [
    ["분석 연도", "독립 변수 (X)", "종속 변수 (Y)", "피어슨 계수 (r)", "피어슨 p-value", "스피어먼 계수 (ρ)", "스피어먼 p-value"],
    ["2020년", "만족도 점수", "자살률", "-0.5036", "0.0393 *", "-0.4020", "1.10e-01"],
    ["", "만족 비율 (%)", "자살률", "-0.5226", "0.0314 *", "-0.3775", "1.35e-01"],
    ["2021년", "만족도 점수", "자살률", "0.0517", "8.44e-01", "0.2466", "3.40e-01"],
    ["", "만족 비율 (%)", "자살률", "-0.1126", "6.67e-01", "0.1683", "5.18e-01"],
    ["2022년", "만족도 점수", "자살률", "0.0785", "7.65e-01", "-0.1128", "6.66e-01"],
    ["", "만족 비율 (%)", "자살률", "-0.0077", "9.77e-01", "-0.1755", "5.01e-01"],
    ["2023년", "만족도 점수", "자살률", "-0.1120", "6.69e-01", "0.0454", "8.63e-01"],
    ["", "만족 비율 (%)", "자살률", "-0.2396", "3.54e-01", "-0.0331", "9.00e-01"],
    ["2024년", "만족도 점수", "자살률", "0.2299", "3.75e-01", "0.2232", "3.89e-01"],
    ["", "만족 비율 (%)", "자살률", "0.0386", "8.83e-01", "0.2074", "4.25e-01"]
]

table3 = doc.add_table(rows=len(yby_corr_data), cols=7)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row_data in enumerate(yby_corr_data):
    for c_idx, val in enumerate(row_data):
        cell = table3.cell(r_idx, c_idx)
        cell.text = val
        is_header = (r_idx == 0)
        bold = is_header or (c_idx == 0 and val != "")
        run = cell.paragraphs[0].runs[0]
        format_run(run, size_pt=9, bold=bold)
        if is_header:
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            run.font.color.rgb = RGBColor(255, 255, 255)
        elif r_idx % 2 == 1:
            set_cell_background(cell, COLOR_LIGHT_GRAY)
        align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER)

p_space2 = doc.add_paragraph()
p_space2.paragraph_format.space_before = Pt(12)

# Insert images for section 4
# Image 3: satisfaction_vs_suicide_scatter.png
img_path3 = os.path.join(project_dir, "images", "satisfaction_vs_suicide_scatter.png")
if os.path.exists(img_path3):
    p_img3 = doc.add_paragraph()
    p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img3.add_run().add_picture(img_path3, width=Inches(5.5))
    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap3.add_run("<그림 3> 삶의 만족도 점수 vs 10만 명당 자살률 산점도 및 선형 추세 (성별/연도 통합)"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap3.paragraph_format.space_after = Pt(12)

# Image 4: correlation_heatmap.png
img_path4 = os.path.join(project_dir, "images", "correlation_heatmap.png")
if os.path.exists(img_path4):
    p_img4 = doc.add_paragraph()
    p_img4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img4.add_run().add_picture(img_path4, width=Inches(5.2))
    p_cap4 = doc.add_paragraph()
    p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap4.add_run("<그림 4> 삶의 세부 만족도 지표(5개 응답) 및 자살률 간의 상관관계 열지도(Heatmap)"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap4.paragraph_format.space_after = Pt(18)


# --- 5. 시도별 자살률과 만족도의 역설 ---
h5 = doc.add_paragraph()
h5.paragraph_format.space_before = Pt(18)
h5.paragraph_format.space_after = Pt(6)
format_run(h5.add_run("5. 시도별 자살률과 만족도의 역설 (The Satisfaction Paradox)"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

p_paradox = doc.add_paragraph()
format_run(p_paradox.add_run(
    "2024년 기준 17개 시도별 삶의 만족도 상세 통계 비율과 자살률 데이터를 대조하면 주관적 감정 만족도와 실제 사망 통계 간의 역설이 뚜렷하게 관찰됩니다.\n"
    "인천이 만족도 최저점 수준이나 자살률은 중간 수준에 머무르는 반면, 충남과 제주는 주민 만족도가 최상위권임에도 자살률은 1, 2위를 차지하고 있습니다."
), size_pt=10)

# Table 4: 2024 Regional Stats (all 17 regions)
regional_header = ["순위", "행정구역", "만족도 점수", "매우 만족 (%)", "약간 만족 (%)", "보통 (%)", "약간 불만족 (%)", "매우 불만족 (%)", "자살률 (명)"]
regional_rows = [
    ["1", "제주특별자치도", "3.44", "12.8", "30.6", "46.4", "8.4", "1.8", "36.3"],
    ["2", "충청남도", "3.56", "16.3", "30.7", "46.7", "4.9", "1.3", "34.8"],
    ["3", "전라남도", "3.34", "10.3", "29.9", "45.4", "12.6", "1.8", "34.5"],
    ["4", "강원특별자치도", "3.42", "12.1", "30.1", "47.7", "8.5", "1.7", "34.3"],
    ["5", "전북특별자치도", "3.38", "10.6", "34.2", "41.0", "11.1", "3.0", "32.3"],
    ["6", "충청북도", "3.39", "13.0", "27.7", "47.0", "10.0", "2.2", "31.8"],
    ["7", "경상북도", "3.35", "9.2", "31.7", "46.2", "10.7", "2.2", "31.6"],
    ["8", "대전광역시", "3.52", "16.0", "31.2", "43.2", "8.1", "1.5", "31.2"],
    ["8", "인천광역시", "3.21", "9.5", "18.5", "58.3", "11.6", "2.2", "31.2"],
    ["10", "부산광역시", "3.32", "12.1", "25.4", "47.5", "12.0", "3.0", "30.3"],
    ["11", "광주광역시", "3.41", "13.2", "27.5", "48.6", "8.9", "1.9", "29.9"],
    ["12", "대구광역시", "3.21", "9.2", "26.9", "42.6", "18.9", "2.5", "29.4"],
    ["13", "울산광역시", "3.43", "13.4", "27.2", "49.6", "8.1", "1.7", "29.2"],
    ["14", "경상남도", "3.40", "10.4", "32.5", "45.2", "10.3", "1.5", "28.5"],
    ["15", "경기도", "3.31", "10.3", "25.4", "51.4", "10.7", "2.3", "28.2"],
    ["16", "서울특별시", "3.44", "12.4", "33.9", "41.8", "9.3", "2.6", "24.1"],
    ["17", "세종특별자치시", "3.35", "16.5", "26.3", "36.4", "17.4", "3.4", "23.0"]
]

table4 = doc.add_table(rows=len(regional_rows)+1, cols=9)
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for c_idx, val in enumerate(regional_header):
    cell = table4.cell(0, c_idx)
    cell.text = val
    run = cell.paragraphs[0].runs[0]
    format_run(run, size_pt=8.5, bold=True)
    set_cell_background(cell, COLOR_PRIMARY_HEX)
    run.font.color.rgb = RGBColor(255, 255, 255)
    align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
for r_idx, row_data in enumerate(regional_rows):
    for c_idx, val in enumerate(row_data):
        cell = table4.cell(r_idx+1, c_idx)
        cell.text = val
        bold = (c_idx in [1, 2, 8])
        run = cell.paragraphs[0].runs[0]
        format_run(run, size_pt=8.5, bold=bold)
        
        # Zebra coloring
        if (r_idx+1) % 2 == 1:
            set_cell_background(cell, COLOR_LIGHT_GRAY)
            
        align_cell(cell, vertical='center', horizontal=WD_ALIGN_PARAGRAPH.CENTER)

p_space3 = doc.add_paragraph()
p_space3.paragraph_format.space_before = Pt(12)

# Insert images for section 5
# Image 5: regional_comparison_bar.png
img_path5 = os.path.join(project_dir, "images", "regional_comparison_bar.png")
if os.path.exists(img_path5):
    p_img5 = doc.add_paragraph()
    p_img5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img5.add_run().add_picture(img_path5, width=Inches(5.5))
    p_cap5 = doc.add_paragraph()
    p_cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap5.add_run("<그림 5> 2024년 시도별 자살률 및 삶의 만족도 비교 (자살률 순 정렬)"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap5.paragraph_format.space_after = Pt(12)

# Image 6: satisfaction_vs_suicide_2024_annotated.png
img_path6 = os.path.join(project_dir, "images", "satisfaction_vs_suicide_2024_annotated.png")
if os.path.exists(img_path6):
    p_img6 = doc.add_paragraph()
    p_img6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img6.add_run().add_picture(img_path6, width=Inches(5.5))
    p_cap6 = doc.add_paragraph()
    p_cap6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(p_cap6.add_run("<그림 6> 2024년 시도별 삶의 만족도 점수 vs 자살률 분포 (시도 산점도)"), size_pt=9, italic=True, color_rgb=COLOR_SECONDARY)
    p_cap6.paragraph_format.space_after = Pt(12)

p_social = doc.add_paragraph()
p_social.paragraph_format.space_before = Pt(8)
format_run(p_social.add_run(
    "역설이 일어나는 요인은 세 가지로 분석됩니다.\n"
    "1) 생태학적 오류 및 고령화 편향: 고령인구 비중이 높은 도(道) 지역은 노인 자살률이 높아 전체 평균 자살률이 올라가는 반면, 설문조사 형태의 주관적 만족도는 일반 인구 분포를 따르므로 취약 노인층의 절망이 과소대표됩니다.\n"
    "2) 사회적 비교 이론: 만족도가 높은 이웃들 사이에 홀로 고독이나 불만을 느끼는 개인은 그 소외감과 심리적 박탈감이 더욱 강하게 나타납니다.\n"
    "3) 보통(무색무취) 비율과 자살률의 연계: 적극적으로 만족/불만족을 표현하지 않고 보통 비율이 높은 '정서적 무관심 지대'가 많은 지역일수록 실제 자살 예방 안전망이 부실하거나 취약 계층의 발굴이 어려워 자살 위험이 상승합니다."
), size_pt=9.5)


# --- 6. 결론 및 제언 ---
h6 = doc.add_paragraph()
h6.paragraph_format.space_before = Pt(18)
h6.paragraph_format.space_after = Pt(6)
format_run(h6.add_run("6. 결론 및 제언"), size_pt=14, bold=True, color_rgb=COLOR_PRIMARY)

conclusions = [
    ("정신건강 정책 지표 다각화: ", "지역 복지 정책 수립 시 단순히 삶의 만족도 서베이 단독 지표만 활용하기보다는 자살 리스크 데이터와 노인 1인가구 등 실질 리스크 인자를 병합하여 위험 지역을 세밀하게 타겟팅해야 합니다."),
    ("무관심 소외 계층(보통 응답층) 발굴: ", "적극적 불만을 노출하지 않고 감정을 유보하거나 '보통'이라고 진술하여 통계적 사각지대에 놓이는 지역 소외 계층을 식별하기 위한 정서 발굴 프로그램 및 탐지망 도입이 필요합니다."),
    ("지방 및 도 지역 특화 복지망 및 성별 예방: ", "자살률이 매우 높게 지속되는 충남, 제주, 전남, 강원 등 도 지역의 취약 가구 밀착 복지 강화와 함께, 남성 만족도 대비 자살률이 2.5배 가량 높은 성별 불균형 문제를 해소하기 위한 남성 특화 프로그램(경제적 실직, 남성 고립 정서 대상) 확대가 강력히 요구됩니다.")
]

for title, body in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run_bold = p.add_run(title)
    format_run(run_bold, size_pt=10, bold=True)
    run_body = p.add_run(body)
    format_run(run_body, size_pt=10)

doc.save(docx_path)
print(f"Word document saved successfully to {docx_path}!")
